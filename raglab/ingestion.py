"""Document loading utilities for RAGLab."""

from __future__ import annotations

from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Any


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".html", ".htm", ".pdf", ".docx"}


@dataclass(frozen=True)
class SourceDocument:
    """Loaded source text plus metadata preserved for retrieval."""

    text: str
    source: str
    metadata: dict[str, Any]


class IngestionError(RuntimeError):
    """Raised when a file cannot be loaded for ingestion."""


def load_documents(file_path: str | Path) -> list[SourceDocument]:
    """Load supported document formats into source documents."""
    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File does not exist: {file_path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise IngestionError(f"Unsupported file type '{suffix}'. Supported: {supported}")

    if suffix == ".txt":
        return [_load_text(path, document_type="txt")]
    if suffix in {".md", ".markdown"}:
        return _load_markdown(path)
    if suffix in {".html", ".htm"}:
        return _load_html(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)

    raise IngestionError(f"Unsupported file type: {suffix}")


def _base_metadata(path: Path, document_type: str) -> dict[str, Any]:
    return {
        "source": path.name,
        "document_type": document_type,
    }


def _load_text(path: Path, document_type: str) -> SourceDocument:
    return SourceDocument(
        text=path.read_text(encoding="utf-8"),
        source=path.name,
        metadata=_base_metadata(path, document_type),
    )


def _load_markdown(path: Path) -> list[SourceDocument]:
    text = path.read_text(encoding="utf-8")
    documents: list[SourceDocument] = []
    current_heading = ""
    current_lines: list[str] = []

    for line in text.splitlines():
        if line.lstrip().startswith("#"):
            if current_lines:
                documents.append(
                    _section_document(path, "markdown", current_heading, current_lines)
                )
                current_lines = []
            current_heading = line.lstrip("#").strip()
        else:
            current_lines.append(line)

    if current_lines or not documents:
        documents.append(
            _section_document(path, "markdown", current_heading, current_lines)
        )

    return [document for document in documents if document.text.strip()]


def _section_document(
    path: Path,
    document_type: str,
    section: str,
    lines: list[str],
) -> SourceDocument:
    metadata = _base_metadata(path, document_type)
    if section:
        metadata["section"] = section
    return SourceDocument(
        text="\n".join(lines).strip(),
        source=path.name,
        metadata=metadata,
    )


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.current_heading = ""
        self.current_tag = ""
        self.current_text: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        self.current_tag = tag.lower()

    def handle_data(self, data: str) -> None:
        text = " ".join(data.split())
        if not text:
            return
        if self.current_tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            self.current_heading = text
        elif self.current_tag in {"p", "li", "td", "th", "div", "span"}:
            self.current_text.append(text)


def _load_html(path: Path) -> list[SourceDocument]:
    parser = _HTMLTextExtractor()
    parser.feed(path.read_text(encoding="utf-8"))
    text = "\n".join(parser.current_text).strip()
    if not text:
        return []
    metadata = _base_metadata(path, "html")
    if parser.current_heading:
        metadata["section"] = parser.current_heading
    return [SourceDocument(text=text, source=path.name, metadata=metadata)]


def _load_pdf(path: Path) -> list[SourceDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise IngestionError(
            "PDF ingestion requires the optional 'pypdf' package."
        ) from exc

    reader = PdfReader(str(path))
    documents: list[SourceDocument] = []
    for index, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        metadata = _base_metadata(path, "pdf")
        metadata["page"] = index
        documents.append(SourceDocument(text=text, source=path.name, metadata=metadata))
    return documents


def _load_docx(path: Path) -> list[SourceDocument]:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestionError(
            "DOCX ingestion requires the optional 'python-docx' package."
        ) from exc

    document = Document(str(path))
    documents: list[SourceDocument] = []
    current_heading = ""
    current_lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if current_lines:
                documents.append(
                    _section_document(path, "docx", current_heading, current_lines)
                )
                current_lines = []
            current_heading = text
        else:
            current_lines.append(text)

    if current_lines:
        documents.append(_section_document(path, "docx", current_heading, current_lines))

    return documents
